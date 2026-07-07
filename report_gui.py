import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import sys
import json
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, date, timezone, timedelta
import re
import textwrap
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import app_config
from widgets import GridTable
from tab_service_record import ServiceRecordMixin
from tab_homepage import HomepageMixin
from tab_call_list import CallListMixin
from tab_liangzhiyibai import LiangzhiyibaiMixin
from tab_live_report import LiveReportMixin
from tab_pre_report import PreReportMixin
from tab_weather_alert import WeatherAlertMixin
from tab_risk_alert import RiskAlertMixin

class ReportApp(ServiceRecordMixin, HomepageMixin, CallListMixin, LiangzhiyibaiMixin, LiveReportMixin, PreReportMixin, WeatherAlertMixin, RiskAlertMixin):
    def __init__(self, root):
        self.root = root
        self.root.title("天气通报生成器")
        # ---------- 设置图标 ----------
        # 获取图标路径（兼容打包后和源码运行）
        if getattr(sys, 'frozen', False):
            base_path = sys._MEIPASS
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))
        icon_path = os.path.join(base_path, "icon.ico")
        if os.path.exists(icon_path):
            self.root.iconbitmap(icon_path)

        self.root.geometry("1280x750")
        self.root.resizable(True, True)

        # ---------- 样式 ----------
        style = ttk.Style()
        style.theme_use("clam")
        style.configure(".", font=("微软雅黑", 12))
        style.configure("TLabel", background="#f0f4f8", foreground="#2c3e50")
        style.configure("TFrame", background="#f0f4f8")
        style.configure("TLabelframe", background="#f0f4f8", foreground="#2c3e50", borderwidth=1, relief="solid")
        style.configure("TLabelframe.Label", font=("微软雅黑", 13, "bold"), background="#f0f4f8", foreground="#1a5276")
        style.configure("TEntry", fieldbackground="white", borderwidth=1, relief="solid")
        style.configure("TCombobox", fieldbackground="white", borderwidth=1, relief="solid")
        style.configure("TButton", background="#2980b9", foreground="white", borderwidth=0, focuscolor="none")
        style.map("TButton", background=[("active", "#3498db")])
        style.configure("Accent.TButton", background="#27ae60", foreground="white")
        style.map("Accent.TButton", background=[("active", "#2ecc71")])

        # ---------- 路径配置 ----------
        if getattr(sys, 'frozen', False):
            # 打包后的临时目录
            self.app_dir = os.path.dirname(sys.executable)
        else:
            self.app_dir = os.path.dirname(os.path.abspath(__file__))
        self.stations_file = os.path.join(self.app_dir, "stations.json")
        self.config_file = os.path.join(self.app_dir, "config.json")

        # ---------- 站点数据 ----------
        self.default_stations = [
            "龙港芦浦华中", "龙港白沙河东", "龙港舥艚中段",
            "云岩瑞联", "芦浦东门垟", "云岩鲸头"
        ]
        self.stations = self.load_stations()

        # ---------- 保存目录（记忆上次，两个标签页独立） ----------
        saved = self.load_config()
        pre_dir = saved.get("pre_save_dir", "")
        live_dir = saved.get("live_save_dir", "")
        weather_dir = saved.get("weather_save_dir", "")
        cwd = os.getcwd()
        self.pre_save_dir = tk.StringVar(value=pre_dir if pre_dir and os.path.isdir(pre_dir) else cwd)
        self.live_save_dir = tk.StringVar(value=live_dir if live_dir and os.path.isdir(live_dir) else cwd)
        self.weather_save_dir = tk.StringVar(value=weather_dir if weather_dir and os.path.isdir(weather_dir) else cwd)
        risk_dir = saved.get("risk_save_dir", "")
        self.risk_save_dir = tk.StringVar(value=risk_dir if risk_dir and os.path.isdir(risk_dir) else cwd)
        self.risk_issue_number = saved.get("risk_issue_number", 1)
        lzy_dir = saved.get("lzy_save_dir", "")
        self.lzy_save_dir = tk.StringVar(value=lzy_dir if lzy_dir else cwd)

        # ---------- 主布局 ----------
        main_panel = ttk.Frame(root)
        main_panel.pack(fill="both", expand=True, padx=10, pady=10)

        # ---------- 右侧站点管理面板 ----------
        self.right_frame = ttk.Frame(main_panel, padding=10, relief="solid", borderwidth=1)
        self.right_frame.pack(side="right", fill="y", padx=(10, 0))
        self.build_station_panel(self.right_frame)

        # ---------- 所有保存栏（先创建，避免标签页切换事件中引用未创建对象） ----------
        self.save_bar_pre = ttk.Frame(root)
        ttk.Label(self.save_bar_pre, text="保存目录：").pack(side="left")
        self.entry_save_pre = ttk.Entry(self.save_bar_pre, textvariable=self.pre_save_dir, width=32)
        self.entry_save_pre.pack(side="left", padx=5)
        ttk.Button(self.save_bar_pre, text="浏览...", command=self.browse_folder_pre).pack(side="left", padx=3)
        ttk.Button(self.save_bar_pre, text="✨ 生成 TXT 文件", command=self.generate_pre_report, style="Accent.TButton").pack(side="right", padx=5)

        self.save_bar_weather = ttk.Frame(root)
        ttk.Label(self.save_bar_weather, text="保存目录：").pack(side="left")
        self.entry_save_weather = ttk.Entry(self.save_bar_weather, textvariable=self.weather_save_dir, width=32)
        self.entry_save_weather.pack(side="left", padx=5)
        ttk.Button(self.save_bar_weather, text="浏览...", command=self.browse_folder_weather).pack(side="left", padx=3)
        ttk.Button(self.save_bar_weather, text="✨ 生成 TXT 文件", command=self.generate_weather_alert, style="Accent.TButton").pack(side="right", padx=5)

        self.save_bar_risk = ttk.Frame(root)
        ttk.Label(self.save_bar_risk, text="保存目录：").pack(side="left")
        self.entry_save_risk = ttk.Entry(self.save_bar_risk, textvariable=self.risk_save_dir, width=32)
        self.entry_save_risk.pack(side="left", padx=5)
        ttk.Button(self.save_bar_risk, text="浏览...", command=self.browse_folder_risk).pack(side="left", padx=3)
        ttk.Button(self.save_bar_risk, text="📄 生成 Word 文档", command=self.generate_risk_alert, style="Accent.TButton").pack(side="right", padx=5)

        self.save_bar_live = ttk.Frame(root)
        ttk.Label(self.save_bar_live, text="保存目录：").pack(side="left")
        self.entry_save_live = ttk.Entry(self.save_bar_live, textvariable=self.live_save_dir, width=32)
        self.entry_save_live.pack(side="left", padx=5)
        ttk.Button(self.save_bar_live, text="浏览...", command=self.browse_folder_live).pack(side="left", padx=3)
        ttk.Button(self.save_bar_live, text="✨ 生成 TXT 文件", command=self.generate_live_report, style="Accent.TButton").pack(side="right", padx=5)

        self.save_bar_lzy = ttk.Frame(root)
        ttk.Label(self.save_bar_lzy, text="保存目录：").pack(side="left")
        self.entry_save_lzy = ttk.Entry(self.save_bar_lzy, textvariable=self.lzy_save_dir, width=32)
        self.entry_save_lzy.pack(side="left", padx=5)
        ttk.Button(self.save_bar_lzy, text="浏览...", command=self._browse_lzy_dir).pack(side="left", padx=3)
        ttk.Button(self.save_bar_lzy, text="📊 写入 Excel", command=self._lzy_write_excel,
                   style="Accent.TButton").pack(side="right", padx=5)

        # ---------- 标签页（Notebook） ----------
        self.notebook = ttk.Notebook(main_panel)
        self.notebook.pack(side="left", fill="both", expand=True)

        # 首页标签页（纯展示文档）
        tab_home = ttk.Frame(self.notebook)
        self.notebook.add(tab_home, text="📖 首页")
        self.build_homepage(tab_home)

        # 叫应名单标签页（表格展示）
        tab_call = ttk.Frame(self.notebook)
        self.notebook.add(tab_call, text="📞 叫应名单")
        self.build_call_list(tab_call)

        # 天气提醒标签页（第 3 项）
        tab_weather = ttk.Frame(self.notebook)
        self.notebook.add(tab_weather, text="🌤天气提醒")
        self.weather_canvas, self.weather_scroll_frame = self._make_scrollable(tab_weather)
        self.build_weather_alert_form()

        # 气象灾害风险提示单标签页（第 4 项）
        tab_risk = ttk.Frame(self.notebook)
        self.notebook.add(tab_risk, text="⚠气象灾害风险提示单")
        self.risk_canvas, self.risk_scroll_frame = self._make_scrollable(tab_risk)
        self.build_risk_alert_form()

        # 预通报标签页（第 5 项）
        tab_pre = ttk.Frame(self.notebook)
        self.notebook.add(tab_pre, text="📋 预通报")
        self.pre_canvas, self.pre_scroll_frame = self._make_scrollable(tab_pre)
        self.build_pre_report_form()

        # 实况通报标签页（第 6 项）
        tab_live = ttk.Frame(self.notebook)
        self.notebook.add(tab_live, text="📊 实况通报")
        self.live_canvas, self.live_scroll_frame = self._make_scrollable(tab_live)
        self.build_live_report_form()

        # 两直一白标签页（第 7 项）
        tab_lzy = ttk.Frame(self.notebook)
        self.notebook.add(tab_lzy, text="📋 两直一白")
        self.lzy_canvas, self.lzy_scroll_frame = self._make_scrollable(tab_lzy)
        self.build_liangzhiyibai()

        # 服务记录标签页（可编辑表格，第 8 项）
        tab_svc = ttk.Frame(self.notebook)
        self.notebook.add(tab_svc, text="📝 服务记录")
        self.build_service_record(tab_svc)

        # 标签页切换绑定（放在所有标签页和保存栏创建之后）
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)

        # ---------- 全局鼠标滚轮处理 ----------
        self._setup_global_wheel()

    def _setup_global_wheel(self):
        """统一全局鼠标滚轮：非编辑状态时页面自由滚动"""
        self._wheel_canvases = {}  # tab_index -> canvas
        # 各标签页对应的 canvas（由 build 方法填充）
        # 在首次 tab 切换时延迟构建映射
        self._wheel_map_built = False

        def _build_wheel_map():
            """构建 tab_index -> canvas 映射"""
            self._wheel_canvases.clear()
            for i in range(self.notebook.index("end")):
                tab_text = self.notebook.tab(i, "text")
                if "首页" in tab_text and hasattr(self, 'home_canvas'):
                    self._wheel_canvases[i] = self.home_canvas
                elif "叫应名单" in tab_text and hasattr(self, 'call_canvas'):
                    self._wheel_canvases[i] = self.call_canvas
                elif "天气提醒" in tab_text and hasattr(self, 'weather_canvas'):
                    self._wheel_canvases[i] = self.weather_canvas
                elif "气象灾害" in tab_text and hasattr(self, 'risk_canvas'):
                    self._wheel_canvases[i] = self.risk_canvas
                elif "预通报" in tab_text and hasattr(self, 'pre_canvas'):
                    self._wheel_canvases[i] = self.pre_canvas
                elif "实况通报" in tab_text and hasattr(self, 'live_canvas'):
                    self._wheel_canvases[i] = self.live_canvas
                elif "两直一白" in tab_text and hasattr(self, 'lzy_canvas'):
                    self._wheel_canvases[i] = self.lzy_canvas
            self._wheel_map_built = True

        def _global_wheel(event):
            if not self._wheel_map_built:
                _build_wheel_map()
            tab_idx = self.notebook.index("current")
            canvas = self._wheel_canvases.get(tab_idx)
            if canvas is None:
                return
            if not canvas.winfo_exists():
                return
            # 检查鼠标是否在该 canvas 区域内
            cx0 = canvas.winfo_rootx()
            cy0 = canvas.winfo_rooty()
            cx1 = cx0 + canvas.winfo_width()
            cy1 = cy0 + canvas.winfo_height()
            if not (cx0 <= event.x_root <= cx1 and cy0 <= event.y_root <= cy1):
                return
            # 非编辑状态下页面自由滚动（文本控件已自适应高度，无需拦截）
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        self.root.bind_all("<MouseWheel>", _global_wheel)

    # ========== 配置持久化 ==========
    # ========== 配置持久化（委托到 app_config 模块） ==========
    def load_config(self):
        return app_config.load_config(self)

    def save_config(self):
        app_config.save_config(self)

    # ========== 站点持久化 ==========
    def load_stations(self):
        return app_config.load_stations(self)

    def save_stations(self, stations_list=None):
        app_config.save_stations(self, stations_list)

    # ========== 站点管理面板 ==========
    def build_station_panel(self, parent):
        ttk.Label(parent, text="📌 站点管理", font=("微软雅黑", 14, "bold")).pack(anchor="w", pady=(0, 10))
        self.station_listbox = tk.Listbox(parent, height=12, width=22, font=("微软雅黑", 12),
                                          selectmode=tk.SINGLE, exportselection=False)
        self.station_listbox.pack(fill="both", expand=True, pady=(0, 10))
        self.refresh_station_listbox()

        add_frame = ttk.Frame(parent)
        add_frame.pack(fill="x", pady=(0, 5))
        self.new_station_entry = ttk.Entry(add_frame, font=("微软雅黑", 12))
        self.new_station_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))
        ttk.Button(add_frame, text="添加", command=self.add_station).pack(side="right")
        ttk.Button(parent, text="删除选中站点", command=self.delete_station).pack(fill="x", pady=(0, 10))

        ttk.Separator(parent, orient="horizontal").pack(fill="x", pady=5)
        ttk.Label(parent, text="💡 输入关键字即可筛选站点\n（自动保存）", font=("微软雅黑", 10), foreground="gray").pack()

    def refresh_station_listbox(self):
        self.station_listbox.delete(0, tk.END)
        for s in self.stations:
            self.station_listbox.insert(tk.END, s)

    def add_station(self):
        name = self.new_station_entry.get().strip()
        if not name:
            messagebox.showwarning("提示", "请输入站点名称")
            return
        if name in self.stations:
            messagebox.showwarning("提示", "站点已存在")
            return
        self.stations.append(name)
        self.refresh_station_listbox()
        self.new_station_entry.delete(0, tk.END)
        self.update_all_station_combos()
        self.save_stations()

    def delete_station(self):
        sel = self.station_listbox.curselection()
        if not sel:
            messagebox.showwarning("提示", "请先选择要删除的站点")
            return
        idx = sel[0]
        del self.stations[idx]
        self.refresh_station_listbox()
        self.update_all_station_combos()
        self.save_stations()

    def update_all_station_combos(self):
        for cb in self.station_combos:
            cb["values"] = self.stations

    # ========== 智能搜索（自定义弹窗版 — 不抢焦点，弹出后仍可编辑） ==========
    def setup_autocomplete(self, combobox, all_values=None):
        """用自定义 Toplevel+Listbox 替代原生下拉框。
           all_values: 完整选项列表，默认 self.stations。"""

        if all_values is None:
            all_values = self.stations
        combobox._all_values = all_values     # 记在控件上，过滤用

        # ── 创建自定义弹出窗 ──
        popup = tk.Toplevel(combobox)
        popup.withdraw()
        popup.overrideredirect(True)        # 无标题栏
        popup.attributes('-topmost', True)  # 置顶
        popup.configure(bg="#3498db")       # 外框颜色 = 蓝色边框

        inner = tk.Frame(popup, bg="white")
        inner.pack(padx=1, pady=1)          # 1px 留白即蓝色边框

        listbox = tk.Listbox(inner, font=("微软雅黑", 12),
                             selectmode=tk.SINGLE, exportselection=False,
                             bg="white", fg="#2c3e50",
                             selectbackground="#3498db",
                             selectforeground="white",
                             relief="flat", borderwidth=0,
                             highlightthickness=0)
        listbox.pack(fill="both", expand=True)

        # ── 定时器 ──
        combobox._autocomplete_timer = None

        # ── 弹窗显示/隐藏 ──
        def _show_popup():
            values = combobox['values']
            if not values:
                _hide_popup()
                return
            listbox.delete(0, tk.END)
            for v in values:
                listbox.insert(tk.END, v)
            n = min(len(values), 8)
            listbox.config(height=n)
            popup.update_idletasks()
            x = combobox.winfo_rootx()
            y = combobox.winfo_rooty() + combobox.winfo_height()
            w = combobox.winfo_width()
            popup.geometry(f"{w}x{listbox.winfo_reqheight() + 2}+{x}+{y}")
            popup.deiconify()

        def _hide_popup():
            popup.withdraw()

        # ── 点击弹窗以外任意位置 → 自动消失 ──
        def _on_global_click(event):
            if not popup.winfo_viewable():
                return
            # 沿控件树向上查找，判断点击是否发生在此 combobox 上
            w = event.widget
            while w is not None and w != self.root:
                if w == combobox:
                    return   # 点击 combobox 本身 → 不隐藏
                try:
                    w = w.master
                except Exception:
                    break
            _hide_popup()

        self.root.bind('<Button-1>', _on_global_click, add='+')

        def _check_and_hide():
            """失焦后延迟检查：鼠标在弹窗上则不隐藏"""
            if not popup.winfo_viewable():
                return
            px, py = popup.winfo_pointerxy()
            x0, y0 = popup.winfo_rootx(), popup.winfo_rooty()
            x1, y1 = x0 + popup.winfo_width(), y0 + popup.winfo_height()
            if not (x0 <= px <= x1 and y0 <= py <= y1):
                _hide_popup()

        # ── 列表项点击 / 回车 → 填入并关闭 ──
        def _on_list_select(event):
            if listbox.curselection():
                text = listbox.get(listbox.curselection()[0])
                combobox.set(text)
                _hide_popup()
                combobox.focus_set()
                combobox.icursor(len(text))

        listbox.bind('<ButtonRelease-1>', _on_list_select)
        listbox.bind('<Return>', _on_list_select)

        # ── 输入框失焦 → 延迟隐藏弹窗 ──
        combobox.bind('<FocusOut>', lambda e: combobox.after(150, _check_and_hide))

        # ── 鼠标点击处理 ──
        def on_click(event):
            w = event.widget.winfo_width()
            if event.x > w - 20:
                # 点击箭头区域 → 显示完整列表（阻止原生下拉框）
                combobox['values'] = list(combobox._all_values)
                _show_popup()
                return 'break'
            # 点击输入区域 → 正常编辑

        combobox.bind('<Button-1>', on_click)

        # ── 键盘处理 ──
        def on_keyrelease(event):
            if event.keysym in ("Up", "Down", "Left", "Right",
                                "Return", "Tab", "Escape",
                                "Control_L", "Control_R",
                                "Shift_L", "Shift_R",
                                "Alt_L", "Alt_R",
                                "Home", "End", "Prior", "Next",
                                "Caps_Lock", "Num_Lock"):
                if event.keysym == "Escape":
                    _hide_popup()
                elif event.keysym == "Down" and popup.winfo_viewable():
                    listbox.focus_set()
                    listbox.selection_set(0)
                elif event.keysym == "Return" and popup.winfo_viewable():
                    _on_list_select(None)
                return

            # 立刻过滤选项
            value = event.widget.get()
            if value == '':
                combobox['values'] = list(combobox._all_values)
            else:
                data = [s for s in combobox._all_values if value.lower() in s.lower()]
                combobox['values'] = data

            # 防抖 250ms 后弹出自定义弹窗（不打断 IME 组字）
            if combobox._autocomplete_timer is not None:
                combobox.after_cancel(combobox._autocomplete_timer)
            combobox._autocomplete_timer = combobox.after(250, _show_popup)

        combobox.bind('<KeyRelease>', on_keyrelease)

    # ========== 滚动容器 ==========
    def _make_scrollable(self, parent):
        """创建带滚动条的容器，返回 (canvas, inner_frame) 供填充"""
        canvas = tk.Canvas(parent, highlightthickness=0, bg="#f0f4f8")
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        inner = ttk.Frame(canvas, padding=10)

        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        canvas_window = canvas.create_window((0, 0), window=inner, anchor="nw")

        def _on_canvas_configure(event):
            canvas.itemconfig(canvas_window, width=event.width)
        canvas.bind("<Configure>", _on_canvas_configure)

        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # 滚轮由全局处理器 _global_wheel_handler 统一管理
        return canvas, inner

    # ========== 首页（展示文档） ==========
    def on_tab_changed(self, event):
        try:
            self._do_tab_changed(event)
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("标签页切换错误", f"{e}")

    def _do_tab_changed(self, event):
        tab_id = self.notebook.select()
        tab_text = self.notebook.tab(tab_id, "text")
        # 所有保存栏已嵌入各标签页内部，root 级栏不再使用
        self.right_frame.pack_forget()
        if "实况通报" in tab_text:
            self.right_frame.pack(side="right", fill="y", padx=(10, 0))

    # ========== 预通报 TXT 生成 ==========
    def generate_pre_report(self):
        try:
            title1 = self.pre_title1.get().strip()
            title2 = self.pre_title2.get().strip()
            system = self.pre_system.get().strip()
            area = self.pre_area.get().strip()
            weather = self.pre_weather.get().strip()
            org1 = self.pre_org1.get().strip()
            org2 = self.pre_org2.get().strip()
            org3 = self.pre_org3.get().strip()
            warn1 = self.pre_warn1.get().strip()
            warn2 = self.pre_warn2.get().strip()
            warn3 = self.pre_warn3.get().strip()
            forecast = self.pre_forecast.get().strip()
            wind = self.pre_wind.get().strip()

            head = f"【{title1}:{title2}强天气预通报】"

            # 构建预警部分（为空自动跳过）
            warn_parts = []
            if org1 and warn1:
                warn_parts.append(f"{org1}已发布{warn1}预警")
            if org2 and warn2:
                warn_parts.append(f"{org2}发布{warn2}预警")
            if org3 and warn3:
                warn_parts.append(f"{org3}发布{warn3}预警")

            if warn_parts:
                p1 = (f"受{system}影响，目前{area}等周边地区出现{weather}天气，"
                      f"{'，'.join(warn_parts)}。")
            else:
                p1 = f"受{system}影响，目前{area}等周边地区出现{weather}天气。"

            p2 = (f"预计{forecast}，有雷雨时可伴有短时强降水、雷电、{wind}、"
                  f"冰雹等强对流天气，请各有关方面注意加强防范。")

            full_text = head + p1 + p2
            wrapped = textwrap.fill(full_text, width=70, break_long_words=False, replace_whitespace=False)

            now_str = datetime.now().strftime("%Y%m%d%H%M")
            filename = f"{now_str}预通报.txt"
            save_dir = self.pre_save_dir.get()
            save_path = os.path.join(save_dir, filename)

            if os.path.exists(save_path):
                base, ext = os.path.splitext(filename)
                count = 2
                while True:
                    new_name = f"{base}_{count}{ext}"
                    save_path = os.path.join(save_dir, new_name)
                    if not os.path.exists(save_path):
                        break
                    count += 1

            with open(save_path, "w", encoding="utf-8") as f:
                f.write(wrapped)

            # 尝试打开文件所在目录（网络路径可能失败，不影响主流程）
            try:
                os.startfile(os.path.normpath(os.path.dirname(save_path)))
            except Exception:
                pass

            # 自动填入服务记录表格
            try:
                self.add_service_row(
                    time_val=f"{title1}:{title2}",
                    item_val="强天气预通报",
                    content_val=full_text,
                )
            except Exception:
                pass

            messagebox.showinfo("成功", f"文件已生成：\n{save_path}")

        except Exception as e:
            messagebox.showerror("错误", f"生成失败：{e}")

    def browse_folder_pre(self):
        folder = filedialog.askdirectory(title="选择预通报保存目录")
        if folder:
            self.pre_save_dir.set(folder)
            self.save_config()

    def browse_folder_live(self):
        folder = filedialog.askdirectory(title="选择实况通报保存目录")
        if folder:
            self.live_save_dir.set(folder)
            self.save_config()

    def browse_folder_risk(self):
        folder = filedialog.askdirectory(title="选择风险提示单保存目录")
        if folder:
            self.risk_save_dir.set(folder)
            self.save_config()

    def browse_folder_weather(self):
        folder = filedialog.askdirectory(title="选择天气提醒保存目录")
        if folder:
            self.weather_save_dir.set(folder)
            self.save_config()

    # ========== 天气提醒 TXT 生成 ==========
    def generate_weather_alert(self):
        try:
            day1 = self.weather_day1.get().strip()
            day2 = self.weather_day2.get().strip()
            fc_day1 = self.weather_fc_day1.get().strip()
            fc_day2 = self.weather_fc_day2.get().strip()
            weather_type = self.weather_type.get().strip()
            rain = self.weather_rain.get("1.0", "end-1c").strip()
            focus = self.weather_focus.get("1.0", "end-1c").strip()

            head = f"【{day1}日-{day2}日龙港天气提醒】"
            p1 = f"预计未来{fc_day1}-{fc_day2}天龙港市{weather_type}。"
            p2 = f"【降水】{rain}"
            p3 = f"【重点关注】{focus}"

            full_text = head + p1 + "\n" + p2 + "\n" + p3
            wrapped = textwrap.fill(full_text, width=70, break_long_words=False, replace_whitespace=False)

            now_str = datetime.now().strftime("%Y%m%d%H%M")
            filename = f"{now_str}天气提醒.txt"
            save_dir = self.weather_save_dir.get()
            save_path = os.path.join(save_dir, filename)

            if os.path.exists(save_path):
                base, ext = os.path.splitext(filename)
                count = 2
                while True:
                    new_name = f"{base}_{count}{ext}"
                    save_path = os.path.join(save_dir, new_name)
                    if not os.path.exists(save_path):
                        break
                    count += 1

            with open(save_path, "w", encoding="utf-8") as f:
                f.write(wrapped)

            # 尝试打开文件所在目录（网络路径可能失败，不影响主流程）
            try:
                os.startfile(os.path.normpath(os.path.dirname(save_path)))
            except Exception:
                pass

            # 自动填入服务记录表格
            try:
                self.add_service_row(
                    time_val=f"{day1}日-{day2}日",
                    item_val="天气提醒",
                    content_val=full_text,
                )
            except Exception:
                pass

            messagebox.showinfo("成功", f"文件已生成：\n{save_path}")

        except Exception as e:
            messagebox.showerror("错误", f"生成失败：{e}")

    # ========== 气象灾害风险提示单 Word 生成 ==========
    def _set_cell_text(self, cell, text, font_name='宋体', size=12, bold=False, align='center'):
        """设置表格单元格文本（含东亚字体支持）"""
        # 清除现有段落
        for p in cell.paragraphs:
            for r in p.runs:
                r._element.getparent().remove(r._element)
        para = cell.paragraphs[0]
        if align == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(str(text))
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        run.bold = bold
        # 垂直居中
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    def _set_warning_cell(self, cell, text, font_name='宋体_GB2312', size=10.5):
        """设置预警风险可能性单元格——按颜色渲染，每种预警换一行"""
        # 清除现有段落
        for p in cell.paragraphs:
            for r in p.runs:
                r._element.getparent().remove(r._element)

        items = [s.strip() for s in text.split("、") if s.strip()]
        if not items:
            return

        # 颜色映射
        def _get_color(item):
            if "红色" in item:
                return "FF6666"
            elif "橙色" in item:
                return "FFA500"
            elif "黄色" in item:
                return "FFFF00"
            return None

        for i, item in enumerate(items):
            if i == 0:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 段落间距收紧
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)

            run = para.add_run(item)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(size)
            run.bold = True

            color = _get_color(item)
            if color:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), color)
                run._element.rPr.append(shd)

        # 垂直居中
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    def generate_risk_alert(self):
        try:
            bj_tz = timezone(timedelta(hours=8))
            bj_now = datetime.now(bj_tz)

            # ---- 收集数据 ----
            pub_unit = self.risk_pub_unit.get().strip()
            pub_time = bj_now.strftime("%Y年%m月%d日%H时")
            body_text = self.risk_body_text.get("1.0", "end-1c").strip()
            current_year = str(bj_now.year)

            # 期数自增
            try:
                user_num = int(self.risk_issue_num.get().strip())
            except ValueError:
                user_num = self.risk_issue_number
            issue_num_to_use = user_num
            self.risk_issue_number = user_num + 1
            self.risk_issue_num.delete(0, tk.END)
            self.risk_issue_num.insert(0, str(self.risk_issue_number))
            self.save_config()

            # 表1 数据
            t1_month = self.risk_t1_month.get().strip()
            t1_day = self.risk_t1_day.get().strip()
            t1_date = f"{t1_month}月{t1_day}日" if t1_month and t1_day else ""
            t1_high = self.risk_t1_high.get().strip()
            t1_mid = self.risk_t1_mid.get().strip()
            t1_low = self.risk_t1_low.get().strip()
            t1_period = self.risk_t1_period.get().strip()
            t1_area = self.risk_t1_area.get().strip()
            t1_disaster = self.risk_t1_disaster.get().strip()
            t1_impact = self.risk_t1_weather_impact.get().strip()

            # 表2 数据（过滤全空行）
            t2_data = []
            for row in self.risk_t2_rows:
                r_month = row["month"].get().strip()
                r_day = row["day"].get().strip()
                r_date = f"{r_month}月{r_day}日" if r_month and r_day else ""
                r_disaster = row["disaster"].get().strip()
                r_area = row["area"].get().strip()
                r_type = row["disaster_type"].get().strip()
                r_cause = row["cause"].get().strip()
                if any([r_disaster, r_area, r_type, r_cause]):
                    t2_data.append({
                        "date": r_date, "disaster": r_disaster,
                        "area": r_area, "disaster_type": r_type, "cause": r_cause,
                    })

            # 落款
            maker = self.risk_maker.get().strip()
            reviewer = self.risk_reviewer.get().strip()
            approver = self.risk_approver.get().strip()

            # ---- 构建 DOCX ----
            doc = Document()

            # 页面设置
            section = doc.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

            STYLE_HEI = '黑体'
            STYLE_SONG = '宋体_GB2312'

            # 标题
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("气象灾害风险提示单")
            r_title.font.name = STYLE_HEI
            r_title._element.rPr.rFonts.set(qn('w:eastAsia'), STYLE_HEI)
            r_title.font.size = Pt(24)
            p_title.paragraph_format.space_after = Pt(6)

            # 期数
            p_num = doc.add_paragraph()
            p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_num = p_num.add_run(f"{current_year}年第{issue_num_to_use}期")
            r_num.font.name = STYLE_HEI
            r_num._element.rPr.rFonts.set(qn('w:eastAsia'), STYLE_HEI)
            r_num.font.size = Pt(12)
            p_num.paragraph_format.space_after = Pt(12)

            # 发布信息行（带下边框）
            p_info = doc.add_paragraph()
            r_info = p_info.add_run(f"发布单位：{pub_unit}              发布时间：{pub_time}")
            r_info.font.name = STYLE_HEI
            r_info._element.rPr.rFonts.set(qn('w:eastAsia'), STYLE_HEI)
            r_info.font.size = Pt(12)
            pPr_info = p_info._element.get_or_add_pPr()
            pBdr_info = OxmlElement('w:pBdr')
            bottom_info = OxmlElement('w:bottom')
            bottom_info.set(qn('w:val'), 'single')
            bottom_info.set(qn('w:sz'), '12')
            bottom_info.set(qn('w:space'), '4')
            bottom_info.set(qn('w:color'), '000000')
            pBdr_info.append(bottom_info)
            pPr_info.append(pBdr_info)
            p_info.paragraph_format.space_after = Pt(12)

            # 正文
            p_body = doc.add_paragraph()
            full_body = f"预计{body_text}" if body_text else "预计"
            r_body = p_body.add_run(full_body)
            r_body.font.name = STYLE_SONG
            r_body._element.rPr.rFonts.set(qn('w:eastAsia'), STYLE_SONG)
            r_body.font.size = Pt(12)
            r_body.bold = True
            p_body.paragraph_format.first_line_indent = Pt(24)
            p_body.paragraph_format.space_after = Pt(12)

            # ---- 表1 ----
            table1 = doc.add_table(rows=3, cols=8)
            table1.style = 'Table Grid'
            # 设置列宽 (8 列，参考模板比例)
            col_widths_t1 = [Cm(2.0), Cm(1.2), Cm(1.3), Cm(1.4), Cm(1.5), Cm(1.4), Cm(2.4), Cm(2.9)]
            for ci, w in enumerate(col_widths_t1):
                for row in table1.rows:
                    row.cells[ci].width = w

            # 合并单元格
            table1.cell(0, 1).merge(table1.cell(0, 3))  # 预警风险可能性 跨 3 列
            table1.cell(0, 0).merge(table1.cell(1, 0))  # 日期 垂直合并
            table1.cell(0, 4).merge(table1.cell(1, 4))  # 影响时段 垂直合并
            table1.cell(0, 5).merge(table1.cell(1, 5))  # 关注区域 垂直合并
            table1.cell(0, 6).merge(table1.cell(1, 6))  # 可能出现灾害种类 垂直合并
            table1.cell(0, 7).merge(table1.cell(1, 7))  # 天气影响 垂直合并

            # 表头文字
            self._set_cell_text(table1.cell(0, 0), "日期", STYLE_HEI, 12)
            self._set_cell_text(table1.cell(0, 1), "预警风险可能性", STYLE_HEI, 12)
            self._set_cell_text(table1.cell(0, 4), "影响时段", STYLE_HEI, 12)
            self._set_cell_text(table1.cell(0, 5), "关注区域", STYLE_HEI, 12)
            self._set_cell_text(table1.cell(0, 6), "可能出现灾害种类", STYLE_HEI, 12)
            self._set_cell_text(table1.cell(0, 7), "天气影响", STYLE_HEI, 12)

            # 子表头 (高/中/低)
            self._set_cell_text(table1.cell(1, 1), "高", STYLE_HEI, 12)
            self._set_cell_text(table1.cell(1, 2), "中", STYLE_HEI, 12)
            self._set_cell_text(table1.cell(1, 3), "低", STYLE_HEI, 12)

            # 数据行
            self._set_cell_text(table1.cell(2, 0), t1_date, STYLE_SONG, 10.5, bold=True)
            self._set_warning_cell(table1.cell(2, 1), t1_high)
            self._set_warning_cell(table1.cell(2, 2), t1_mid)
            self._set_warning_cell(table1.cell(2, 3), t1_low)
            self._set_cell_text(table1.cell(2, 4), t1_period, STYLE_SONG, 10.5, bold=True)
            self._set_cell_text(table1.cell(2, 5), t1_area, STYLE_SONG, 10.5, bold=True)
            self._set_cell_text(table1.cell(2, 6), t1_disaster, STYLE_SONG, 10.5, bold=True)
            self._set_cell_text(table1.cell(2, 7), t1_impact, STYLE_SONG, 10.5, bold=True)

            # 两张表格之间的空行
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(6)
            spacer.paragraph_format.space_after = Pt(6)

            # ---- 表2（仅非空时生成） ----
            if t2_data:
                n_data = len(t2_data)
                table2 = doc.add_table(rows=2 + n_data, cols=5)
                table2.style = 'Table Grid'
                col_widths_t2 = [Cm(2.3), Cm(2.7), Cm(2.3), Cm(3.9), Cm(2.7)]
                for ci, w in enumerate(col_widths_t2):
                    for row in table2.rows:
                        row.cells[ci].width = w

                headers_t2 = ["日期", "强天气灾害", "关注区域", "可能出现灾害种类", "致灾原因"]
                for ci in range(5):
                    table2.cell(0, ci).merge(table2.cell(1, ci))
                    self._set_cell_text(table2.cell(0, ci), headers_t2[ci], STYLE_HEI, 12)

                for ri, rd in enumerate(t2_data):
                    vals = [rd["date"], rd["disaster"], rd["area"], rd["disaster_type"], rd["cause"]]
                    for ci, val in enumerate(vals):
                        self._set_cell_text(table2.cell(2 + ri, ci), val, STYLE_SONG, 10.5, bold=True)

            # ---- 落款（上黑线 + 签名 + 下黑线） ----
            # 落款前空行（紧缩间距）
            p_spacer = doc.add_paragraph()
            p_spacer.paragraph_format.space_before = Pt(4)
            p_spacer.paragraph_format.space_after = Pt(0)

            # 上黑线段落（无额外间距）
            p_line1 = doc.add_paragraph()
            p_line1.paragraph_format.space_before = Pt(0)
            p_line1.paragraph_format.space_after = Pt(2)
            pPr_l1 = p_line1._element.get_or_add_pPr()
            pBdr_l1 = OxmlElement('w:pBdr')
            top_l1 = OxmlElement('w:top')
            top_l1.set(qn('w:val'), 'single')
            top_l1.set(qn('w:sz'), '12')
            top_l1.set(qn('w:space'), '1')
            top_l1.set(qn('w:color'), '000000')
            pBdr_l1.append(top_l1)
            pPr_l1.append(pBdr_l1)

            # 签名行（紧贴黑线）
            p_sig = doc.add_paragraph()
            p_sig.paragraph_format.space_before = Pt(2)
            p_sig.paragraph_format.space_after = Pt(2)
            sig_parts = []
            if maker:
                sig_parts.append(f"制作：{maker}")
            if reviewer:
                sig_parts.append(f"审核：{reviewer}")
            if approver:
                sig_parts.append(f"签发：{approver}")
            sig_text = "              ".join(sig_parts) if sig_parts else "制作：              审核：              签发："
            r_sig = p_sig.add_run(sig_text)
            r_sig.font.name = STYLE_HEI
            r_sig._element.rPr.rFonts.set(qn('w:eastAsia'), STYLE_HEI)
            r_sig.font.size = Pt(12)

            # 下黑线段落（无额外间距）
            p_line2 = doc.add_paragraph()
            p_line2.paragraph_format.space_before = Pt(0)
            p_line2.paragraph_format.space_after = Pt(0)
            pPr_l2 = p_line2._element.get_or_add_pPr()
            pBdr_l2 = OxmlElement('w:pBdr')
            bottom_l2 = OxmlElement('w:bottom')
            bottom_l2.set(qn('w:val'), 'single')
            bottom_l2.set(qn('w:sz'), '12')
            bottom_l2.set(qn('w:space'), '1')
            bottom_l2.set(qn('w:color'), '000000')
            pBdr_l2.append(bottom_l2)
            pPr_l2.append(pBdr_l2)

            # ---- 保存文件 ----
            now_str = bj_now.strftime("%Y%m%d%H%M")
            filename = f"气象灾害风险提示单（{current_year}年第{issue_num_to_use}期）.docx"
            save_dir = self.risk_save_dir.get()
            save_path = os.path.join(save_dir, filename)

            if os.path.exists(save_path):
                base, ext = os.path.splitext(filename)
                count = 2
                while True:
                    new_name = f"{base}_{count}{ext}"
                    save_path = os.path.join(save_dir, new_name)
                    if not os.path.exists(save_path):
                        break
                    count += 1

            doc.save(save_path)

            # 打开目录
            try:
                os.startfile(os.path.normpath(os.path.dirname(save_path)))
            except Exception:
                pass

            # 自动填入服务记录
            try:
                self.add_service_row(
                    time_val=f"{current_year}年第{issue_num_to_use}期",
                    item_val="气象灾害风险提示单",
                    content_val=f"已生成：{save_path}",
                )
            except Exception:
                pass

            messagebox.showinfo("成功", f"Word 文档已生成：\n{save_path}")

        except Exception as e:
            messagebox.showerror("错误", f"生成失败：{e}")

    def generate_live_report(self):
        try:
            # 优先使用手动编辑框内容
            manual = self.live_manual_text.get("1.0", "end-1c").strip()
            if manual:
                full_text = manual
                wrapped = textwrap.fill(full_text, width=70, break_long_words=False, replace_whitespace=False)
            else:
                title1 = self.entry_title1.get().strip()
                title2 = self.entry_title2.get().strip()
                sys_1h = self.cb_1h_system.get().strip()
                weather_1h = self.cb_1h_weather.get().strip()
                avg_1h = self.entry_1h_avg.get().strip()
                stations_1h = [(e[0].get().strip(), e[1].get().strip()) for e in self.entries_1h_stations]

                avg_3h = self.entry_3h_avg.get().strip()
                stations_3h = [(e[0].get().strip(), e[1].get().strip()) for e in self.entries_3h_stations]
                max_sta_3h = self.cb_3h_max_sta.get().strip()
                max_val_3h = self.entry_3h_max_val.get().strip()
                wind_sta_3h = self.cb_3h_wind_sta.get().strip()
                wind_speed_3h = self.entry_3h_wind_speed.get().strip()
                wind_scale_3h = self.entry_3h_wind_scale.get().strip()

                fc_time = self.entry_fc_time.get().strip()
                fc_weather = self.entry_fc_weather.get().strip()
                fc_rain1 = self.entry_fc_rain1.get().strip()
                fc_rain2 = self.entry_fc_rain2.get().strip()

                head = f"【{title1}:{title2}实况通报】"
                p1 = (f"受{sys_1h}影响，近1小时，龙港市出现{weather_1h}，"
                      f"全市面雨量{avg_1h}毫米；"
                      f"单站总雨量前三分别为{stations_1h[0][0]}{stations_1h[0][1]}毫米、"
                      f"{stations_1h[1][0]}{stations_1h[1][1]}毫米、"
                      f"{stations_1h[2][0]}{stations_1h[2][1]}毫米。")
                part1 = head + p1
                p3 = (f"近3小时，全市面雨量{avg_3h}毫米；"
                      f"单站总雨量前三分别为{stations_3h[0][0]}{stations_3h[0][1]}毫米、"
                      f"{stations_3h[1][0]}{stations_3h[1][1]}毫米、"
                      f"{stations_3h[2][0]}{stations_3h[2][1]}毫米。"
                      f"单站1小时雨量最大出现在{max_sta_3h}{max_val_3h}毫米。"
                      f"风力最大出现在{wind_sta_3h}{wind_speed_3h}米/秒（{wind_scale_3h}级）。")
                p4 = (f"【临近预报】预计未来{fc_time}小时，龙港市仍有{fc_weather}，"
                      f"雨量{fc_rain1}-{fc_rain2}毫米。")

                full_text = part1 + "\n" + p3 + "\n" + p4
                wrapped = textwrap.fill(full_text, width=70, break_long_words=False, replace_whitespace=False)

            now_str = datetime.now().strftime("%Y%m%d%H%M")
            filename = f"{now_str}实况通报.txt"
            save_dir = self.live_save_dir.get()
            save_path = os.path.join(save_dir, filename)

            if os.path.exists(save_path):
                base, ext = os.path.splitext(filename)
                count = 2
                while True:
                    new_name = f"{base}_{count}{ext}"
                    save_path = os.path.join(save_dir, new_name)
                    if not os.path.exists(save_path):
                        break
                    count += 1

            with open(save_path, "w", encoding="utf-8") as f:
                f.write(wrapped)

            # 尝试打开文件所在目录（网络路径可能失败，不影响主流程）
            try:
                os.startfile(os.path.normpath(os.path.dirname(save_path)))
            except Exception:
                pass

            # 自动填入服务记录表格（手动编辑模式跳过）
            if not manual:
                try:
                    self.add_service_row(
                        time_val=f"{title1}:{title2}",
                        item_val="实况通报",
                        content_val=full_text,
                    )
                except Exception:
                    pass

            messagebox.showinfo("成功", f"文件已生成：\n{save_path}")

        except Exception as e:
            messagebox.showerror("错误", f"生成失败：{e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = ReportApp(root)
    root.mainloop()
